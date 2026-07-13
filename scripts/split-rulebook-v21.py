from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table


CHAPTER_TITLES = {
    "0": "规则边界、术语与唯一结算顺序",
    "1": "共同设定与战役开关",
    "2": "核心检定",
    "3": "动作与能力卡",
    "4": "战斗规则",
    "5": "伤害、状态与恢复",
    "5A": "角色创建总则",
    "6": "普通人",
    "7": "魔术师与魔术",
    "8": "御主与契约",
    "9": "从者与职阶",
    "10": "从者技能",
    "11": "宝具",
    "12": "令咒",
    "13": "调查、真名与神秘隐匿",
    "14": "主持人工具",
    "15": "四种战役框架",
    "16": "完整结算范例",
    "A": "桌边速查",
    "B": "角色卡",
    "C": "标准术式书",
    "D": "魔术礼装目录",
    "E": "能力模板与示例敌人",
    "F": "设定来源与版本说明",
}


SPLITS = (
    (
        "玩家核心分册",
        {"0", "1", "2", "3", "4", "5", "5A", "6", "7", "8", "9", "10", "11", "12", "13", "16", "A", "B", "F"},
        "规则、创建、战斗、御主、从者、宝具、调查、范例、速查与角色卡。",
    ),
    (
        "主持人工具与战役框架分册",
        {"0", "1", "2", "3", "4", "5", "13", "14", "15", "16", "E", "F"},
        "主持人运行所需的核心时序、敌人、调查、战役框架、胜负条款与模板。",
    ),
    (
        "角色、术式与礼装资源库",
        {"0", "3", "5A", "6", "7", "8", "9", "10", "11", "12", "13", "C", "D", "E", "F"},
        "角色构筑、能力卡、职阶、宝具、标准术式、礼装与可复用模板。",
    ),
)


def table_values(table: Table) -> list[str]:
    if not table.rows:
        return []
    return [cell.text.strip() for cell in table.rows[0].cells]


def chapter_banner(table: Table) -> str | None:
    if len(table.rows) != 1:
        return None
    values = table_values(table)
    if len(values) < 2:
        return None
    label = values[0]
    title = values[1].splitlines()[0].strip()
    expected = CHAPTER_TITLES.get(label)
    if expected and title.startswith(expected):
        return label
    return None


def toc_token(value: str) -> str | None:
    value = value.strip()
    match = re.fullmatch(r"第\s*([0-9]+A?)\s*章", value)
    if match:
        return match.group(1)
    match = re.fullmatch(r"附录\s*([A-F])", value)
    if match:
        return match.group(1)
    return None


def remove_row(row) -> None:
    row._tr.getparent().remove(row._tr)


def rewrite_front_matter(doc, split_name: str, scope: str, keep: set[str]) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "《零之圣杯》通用圣杯战争规则书":
            paragraph.text = f"《零之圣杯》{split_name}"
        elif text == "稳定公测版 v2.1 · 规则闭环与平衡升级 · 非官方同人桌面角色扮演规则":
            paragraph.text = f"Null Grail Core d20 v2.1 · {scope}"

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if "通用圣杯战争规则书 v2.1" in paragraph.text:
                paragraph.text = paragraph.text.replace("通用圣杯战争规则书 v2.1", f"{split_name} v2.1")

    for table in doc.tables:
        values = table_values(table)
        if values[:3] != ["编号", "内容", "用途"]:
            continue
        for row in list(table.rows[1:]):
            token = toc_token(row.cells[0].text)
            if token and token not in keep:
                remove_row(row)
        break

    doc.core_properties.title = f"《零之圣杯》{split_name} v2.1"
    doc.core_properties.subject = scope
    doc.core_properties.comments = "由 v2.1 完整规则书自动生成的阅读分册；规则效力以同版本完整规则书为准。"


def create_split(source: Path, destination: Path, split_name: str, keep: set[str], scope: str) -> None:
    doc = Document(source)
    current_chapter: str | None = None
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:tbl"):
            token = chapter_banner(Table(child, doc))
            if token:
                current_chapter = token
        if current_chapter is not None and current_chapter not in keep:
            body.remove(child)

    rewrite_front_matter(doc, split_name, scope, keep)
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: split-rulebook-v21.py FULL_RULEBOOK.docx OUTPUT_DIR")
    source = Path(sys.argv[1]).resolve()
    output_dir = Path(sys.argv[2]).resolve()
    for split_name, keep, scope in SPLITS:
        destination = output_dir / f"《零之圣杯》Null Grail Core d20 v2.1_{split_name}.docx"
        create_split(source, destination, split_name, keep, scope)
        print(destination)


if __name__ == "__main__":
    main()
